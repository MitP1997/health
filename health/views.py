from django.shortcuts import render, get_object_or_404, redirect, render_to_response
from django.contrib import auth
from django.contrib.auth.models import User
from django.utils import timezone
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import update_last_login
from django.core.exceptions import ObjectDoesNotExist
from .models import *
from django.core.mail import EmailMessage, send_mail
import csv
import random
import codecs
import urllib
import requests
from django.http import HttpResponse
from xlsxwriter.workbook import Workbook
import json
import os,subprocess
from django.conf import settings
from django.template import RequestContext
from django.core.mail import EmailMessage, send_mail
#import docx
from docx import *
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.enum.table import *
from docx.enum.text import *
import zipfile
import StringIO
from django.db.models import Sum
import random
import operator
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage

# Create your views here.

def PatientRegister(request):
    aadhaarid = request.POST.get('aadhaarid')
    name = request.POST.get('name')
    dob = request.POST.get('dob')
    address = request.POST.get('address')
    gender = request.POST.get('gender')
    contact = request.POST.get('contact')

    height = request.POST.get('height')
    weight = request.POST.get('weight')
    bloodgroup = request.POST.get('bloodgroup')
    medicalconditions = request.POST.get('medicalconditions')
    alergiesandreactions = request.POST.get('alergiesandreactions')
    ongoingmedications = request.POST.get('ongoingmedications')
    familyhistory = request.POST.get('familyhistory')
    emergencycontact = request.POST.get('emergencycontact')

    patient = Patients()
    patient.aadhaarid = aadhaarid
    patient.name = name
    patient.dob = dob
    patient.address = address
    patient.gender = gender
    patient.contact = contact
    patient.height = height
    patient.weight = weight
    patient.bloodgroup = bloodgroup
    patient.medicalconditions = medicalconditions
    patient.alergiesandreactions = alergiesandreactions
    patient.ongoingmedications = ongoingmedications
    patient.familyhistory = familyhistory
    patient.emergencycontact = emergencycontact
    patient.save()
    response_array = {}
    response_array['status'] = 1
    response_array['data'] = 'success'
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")


def FetchPatientDetails(request):
    aadhaarid = request.GET.get('aadhaarid')
    patient = Patients.objects.get(aadhaarid=aadhaarid)
    response_array = {}
    response_array['status'] = 1
    response_array['data'] = {}
    response_array['data']['aadhaarid'] = patient.aadhaarid
    response_array['data']['name'] = patient.name
    response_array['data']['dob'] = patient.dob
    response_array['data']['address'] = patient.address
    response_array['data']['gender'] = patient.gender
    response_array['data']['contact'] = patient.contact
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")


def GetAllPatientRecords(request):
    aadhaarid = request.GET.get('aadhaarid')
    patientRecords = Patientrecords.objects.filter(patient=Patients.objects.get(aadhaarid=aadhaarid))
    response_array = {}
    response_array['status'] = 1
    response_array['data'] = []
    for patientRecord in patientRecords:
        patientRecordObj = {}
        patientRecordObj['symptoms'] = patientRecord.symptoms
        patientRecordObj['diagnosis'] = patientRecord.diagnosis
        patientRecordObj['medication'] = patientRecord.medication
        patientRecordObj['addedby'] = patientRecord.addedby.name
        patientRecordObj['attatchmentlink'] = patientRecord.attatchmentlink
        patientRecordObj['createdat'] = patientRecord.createdat
        patientRecordObj['createdon'] = patientRecord.createdon
        patientRecordObj['id'] = patientRecord.id
        response_array['data'].append(patientRecordObj)
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")

def CreateNewPatientRecord(request):
    aadhaarid = request.POST.get('aadhaarid')
    patient = Patients.objects.get(aadhaarid=aadhaarid)
    symptoms = request.POST.get('symptoms')
    diagnosis = request.POST.get('diagnosis')
    medication = request.POST.get('medication')
    addedby = request.POST.get('addedby')
    attatchmentlink = request.POST.get('attatchmentlink')
    createdat = "11:11"
    createdon = "11-02-2018"
    patientRecord = Patientrecords()
    patientRecord.patient = patient
    patientRecord.symptoms = symptoms
    patientRecord.diagnosis = diagnosis
    patientRecord.medication = medication
    patientRecord.addedby = Doctors.objects.get(id=addedby)
    patientRecord.attatchmentlink = attatchmentlink
    patientRecord.createdat = createdat
    patientRecord.createdon = createdon
    patientRecord.save()

    p=subprocess.Popen("python blockchain.py 1:"+aadhaarid+"2:"+symptoms+"3:"+diagnosis+"4:"+medication ,stdout=subprocess.PIPE,shell=True)
    (output,err)=p.communicate()
    p_status=p.wait()
    print (output)

    response_array = {}
    response_array['status'] = 1
    response_array['data'] = 'success'
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")

##TODO : TEST
def CreateNewVidChatRoom(request):
    aadhaarid = request.POST.get('aadhaarid')
    patient = Patients.objects.get(aadhaarid=aadhaarid)
    createdat = request.POST.get('createdat')
    videorecords = Videorecords()
    videorecords.patient = patient
    videorecords.doctor = Doctors.objects.get(id=request.POST.get('doctorid'))
    videorecords.createdat = createdat
    videorecords.save()

    response_array = {}
    response_array['status'] = 1
    response_array['data'] = 'success'
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")


def SearchDoctor(request):
    name = request.GET.get('name')
    doctor = Doctors.objects.filter(name__icontains=name)

    response_array = {}
    response_array['status'] = 1
    response_array['data'] = []
    for doctor in doctor:
        doctorObj = {}
        doctorObj['available'] = doctor.available
        doctorObj['name'] = doctor.name
        doctorObj['speciality'] = doctor.speciality
        doctorObj['currentlocation'] = doctor.currentlocation.name
        response_array['data'].append(doctorObj)
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")


def SearchDoctorBySpeciality(request):
    speciality = request.GET.get('speciality')
    doctor = Doctors.objects.filter(speciality__icontains=speciality)

    response_array = {}
    response_array['status'] = 1
    response_array['data'] = []
    for doctor in doctor:
        doctorObj = {}
        doctorObj['available'] = doctor.available
        doctorObj['name'] = doctor.name
        doctorObj['speciality'] = doctor.speciality
        doctorObj['currentlocation'] = doctor.currentlocation.name
        response_array['data'].append(doctorObj)
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")


def SearchHospitalByName(request):
    hospital = request.GET.get('hospital')
    hosp = Hospitals.objects.filter(name__icontains=hospital)
    response_array = {}
    response_array['status'] = 1
    response_array['data'] = []
    for hospital in hosp:
        hospitalObj = {}
        hospitalObj['name'] = hospital.name
        hospitalObj['location'] = hospital.location
        hospitalObj['address'] = hospital.address.name
        response_array['data'].append(hospitalObj)
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")

##TODO test
def EmergencyDoctorSearch(request):
    latitude = request.GET.get('latitude')
    longitude = request.GET.get('longitude')
    speciality = request.GET.get('speciality')
    hospitals = Hospitals.objects.all()
    noOfClosestDoctors = 5

    hospitalDistDict = {}
    for hospital in hospitals:
        location = hospital.location
        latitudeDb = float(location.split(',')[0])
        longitudeDb = float(location.split(',')[1])
        sqDist = ((latitudeDb - latitude)**2 +(longitudeDb - longitude)**2)
        hospitalDistDict[hospital.name] = sqDist
    sorted_hospitals = sorted(hospitalDistDict.items(), key=operator.itemgetter(1))[::-1]

    response_array = {}
    response_array['status'] = 1
    response_array['data'] = []
    for i in range(noOfClosestDoctors):
        doctor = Doctors.objects.get(currentlocation=Hospitals.objects.get(name=sorted_hospitals[sorted_hospitals.keys()[i]]))
        doctorObj = {}
        doctorObj['available'] = doctor.available
        doctorObj['name'] = doctor.name
        doctorObj['speciality'] = doctor.speciality
        doctorObj['currentlocation'] = doctor.currentlocation.name
        response_array['data'].append(doctorObj)
    return HttpResponse(json.dumps(response_array), status=200, content_type="application/json")



def DoctorRegister(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        username = request.POST.get('username')
        doctorCount = Doctors.objects.filter(username=username).count()
        if doctorCount != 0:
            return HttpResponse('username already taken')
        password = request.POST.get('password')
        speciality = request.POST.get('speciality')
        doctor = Doctors()
        doctor.name = name
        doctor.username = username
        doctor.password = password
        if speciality is not None:
            doctor.speciality = speciality
        doctor.save()
        return redirect('../DoctorLogin')
    else:
        return render('./Doctors/login.html')

def DoctorLogin(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        doctor = Doctors.objects.get(username=username)
        if doctor is None == 0:
            return render(request,'./Doctors/login.html',{'error':"Username Does Not exist"})
        if password == doctor.password:
            request.session['doctor'] = doctor.username
            return redirect('../DoctorLanding')
        else:
            return render(request,'./Doctors/login.html',{'error':"Invalid Password"})
    else:
        return render(request,'./Doctors/login.html')

def DoctorAddRecord(request):
    if request.method == 'POST':
        aadhaarid = request.POST.get('aadhaarid').replace(" ","")
        patient = Patients.objects.get(aadhaarid=aadhaarid)
        symptoms = request.POST.get('symptoms')
        diagnosis = request.POST.get('diagnosis')
        medication = request.POST.get('medication')
        addedby = request.POST.get('addedby')
        #
        # myfile = request.FILES.get('attachment')
        # fs = FileSystemStorage()
        # filename = fs.save(myfile.name, myfile)
        # uploaded_file_url = fs.url(filename)

        createdat = "01:01"
        createdon = request.POST.get('createdon')
        patientRecord = Patientrecords()
        patientRecord.patient = patient
        patientRecord.symptoms = symptoms
        patientRecord.diagnosis = diagnosis
        patientRecord.medication = medication
        patientRecord.addedby = Doctors.objects.get(username=addedby)
        patientRecord.attatchmentlink = "../"
        patientRecord.createdat = createdat
        patientRecord.createdon = createdon
        patientRecord.save()
        # subprocess.call("python blockchain.py 1:"+aadhaarid+"2:"+symptoms+"3:"+diagnosis+"4:"+medication, shell = True)
        p=subprocess.Popen("python blockchain.py 1:"+aadhaarid+"2:"+symptoms+"3:"+diagnosis+"4:"+medication ,stdout=subprocess.PIPE,shell=True)
        (output,err)=p.communicate()
        p_status=p.wait()
        print (output)

        # return render(request,'./Doctors/AddRecords.html')
        return redirect('../DoctorLanding2/')
    else:
        print(Doctors.objects.get(username=request.session['doctor']))
        return render(request,'./Doctors/AddRecords.html',{'doctorInfo':Doctors.objects.get(username=request.session['doctor'])})

def DoctorLanding(request):
    doctor = request.session['doctor']
    videorecords =  Videorecords.objects.filter(doctor=Doctors.objects.get(username=doctor))
    return render(request,'./Doctors/VideoTable.html',{'videorecords':videorecords})

def DoctorLanding2(request):
    if request.method == "POST":
        doctor = request.session['doctor']
        aadhaarid = request.POST.get("aadhaarid")
        return render(request,'./Doctors/PatientRecords.html',{"patients":Patients.objects.all(), "records":Patientrecords.objects.filter(patient=Patients.objects.get(aadhaarid=aadhaarid))})
    else:
        return render(request,'./Doctors/PatientRecords.html',{"patients":Patients.objects.all()})

def ReportGeneration(request):
    aadhaarid = request.GET.get('aadhaarid')
    patient = Patients.objects.get(aadhaarid=aadhaarid)
    doc = Document()
    doc.add_paragraph('Aadhaar ID: '+aadhaarid+'\nName: '+patient.name+'\nD.O.B.: '+patient.dob+'\nAddress: '+patient.address+'\nGender: '+patient.gender+'\nContact: '+patient.contact+'\nHeight: '+patient.height+'\nWeight: '+patient.weight+'\nBlood Group: '+patient.bloodgroup+'\nMedical Conditions: '+patient.medicalconditions+'\nAlergies and Reactions: '+patient.alergiesandreactions+'\nOngoing Medications: '+patient.ongoingmedications+'\nFamily History: '+patient.familyhistory+'\nEmergency Contact: '+patient.emergencycontact)
    patientrecord = Patientrecords.objects.filter(patient=patient)
    recordTable = doc.add_table(rows = patientrecord.count()+1, cols = 6, style='TableGrid')
    row = 0
    recordTable.cell(row,0).text = 'Symptoms'
    recordTable.cell(row,1).text = 'Diagnosis'
    recordTable.cell(row,2).text = 'Diagnosed by'
    recordTable.cell(row,3).text = 'Medication'
    recordTable.cell(row,4).text = 'Added on'
    recordTable.cell(row,5).text = 'Related Records'
    row += 1
    for record in patientrecord:
        recordTable.cell(row,0).text = ''+record.symptoms
        recordTable.cell(row,1).text = ''+record.diagnosis
        recordTable.cell(row,3).text = ''+record.medication
        addedByName = Doctors.objects.get(id=record.addedby.id)
        recordTable.cell(row,2).text = ''+addedByName.name
        recordTable.cell(row,4).text = ''+record.createdon
        recordTable.cell(row,5).text = ''+record.attatchmentlink
        row += 1
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    doc.save(response)
    return response

def PharmacyRegister(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        username = request.POST.get('username')
        pharmacyCount = Pharmacies.objects.filter(username=username).count()
        if pharmacyCount != 0:
            return render(request,'./Pharmacy/login.html',{"error":"Username already taken!"})
        password = request.POST.get('password')
        address = request.POST.get('address')
        timing = request.POST.get('timing')
        location = request.POST.get('location')
        pharmacy = Pharmacies()
        pharmacy.name = name
        pharmacy.username = username
        pharmacy.password = password
        pharmacy.address = address
        pharmacy.timing = timing
        pharmacy.location = location
        pharmacy.save()
        return redirect('../PharmacyLogin')
    else:
        return render(request,'./Pharmacy/login.html')

def PharmacyLogin(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        pharmacy = Pharmacies.objects.get(username=username)
        if pharmacy is None:
            return render(request,'./Pharmacy/login.html',{'error':"Username does not exists"})
        if password == pharmacy.password:
            request.session['pharmacy'] = pharmacy.username
            return redirect('../PharmacyLanding')
        else:
            return render(request,'./Pharmacy/login.html',{'error':"Invalid Password"})
    else:
        return render(request,'./Pharmacy/login.html')

def PharmacyAddMedicine(request):
    if request.method == 'POST':
        pharmacy = request.session['pharmacy']
        medicinename = request.POST.get('medicinename')
        price = request.POST.get('price')
        medicineCount = Medicines.objects.filter(name=medicinename).count()
        if medicineCount == 0:
            medicine = Medicines()
            medicine.name = medicinename
            medicine.amount = price
            medicine.save()
        quantity = request.POST.get('quantity')
        if (Medicinestocks.objects.filter(medicine=Medicines.objects.get(name=medicinename)).filter(pharmacy=Pharmacies.objects.get(username=pharmacy)).count()==0):
            medicinestocks = Medicinestocks()
            medicinestocks.pharmacy = Pharmacies.objects.get(username=pharmacy)
            medicinestocks.medicine = Medicines.objects.get(name=medicinename)
            medicinestocks.quantity = quantity
            medicinestocks.save()
        else:
            oldQuantity = Medicinestocks.objects.filter(medicine=Medicines.objects.get(name=medicinename)).filter(pharmacy=Pharmacies.objects.get(username=pharmacy))[0].quantity
            newQuantity = str( int(oldQuantity) + int(quantity) )
            if (newQuantity < 0):
                return render(request,'./Pharmacy/Add_meds.html',{'error':"Invalid Quantity Change"})
            else:
                Medicinestocks.objects.filter(medicine=Medicines.objects.get(name=medicinename)).filter(pharmacy=Pharmacies.objects.get(username=pharmacy)).update(quantity=newQuantity)
        return redirect('../PharmacyLanding')
    else:

        return render(request,'./Pharmacy/Add_meds.html')

def PharmacyLanding(request):
    if request.method =='POST':
        pass
    else:
        pharmacy = request.session['pharmacy']
        medicinestocks = Medicinestocks.objects.filter(pharmacy=Pharmacies.objects.get(username=pharmacy))
        return render(request,'./Pharmacy/MedicineTable.html',{'medicines': medicinestocks})

def PharmacyInvoice(request):

    return render(request,'./Pharmacy/invoice.html')

def PharmacyInvoiceGenerator(request):
    if request.method == 'POST':
        count = request.POST.get('count')
        count = int(count)
        itemFetched = []
        total = 0
        for i in range(count):
            medicine = (Medicines.objects.get(id=request.POST.get('medicineid'+str(i+1))))
            qty = (request.POST.get('quantity'+str(i+1)))
            obj = {'name':medicine.name,'amount':medicine.amount,'quantity':qty,'subtotal':'150'}
            total = total + 150
            itemFetched.append(obj)
        return render(request,'./Pharmacy/Invoice.html',{'items':itemFetched, 'total':total})
    else:
        pharmacy = Pharmacies.objects.get(username=request.session['pharmacy'])
        medicine = Medicinestocks.objects.filter(pharmacy=pharmacy)
        return render(request,'./Pharmacy/invoiceform.html',{'medicineList':medicine})

def HospitalRegister(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        location = request.POST.get('location')
        address = request.POST.get('address')
        username = request.POST.get('username')
        hospitalCount = Hospitals.objects.filter(username=username).count()
        if hospitalCount != 0:
            return render(request,'./Equipment/login.html',{"error":"Username already taken!"})
        password = request.POST.get('password')
        hospital = Hospitals()
        hospital.name = name
        hospital.username = username
        hospital.password = password
        hospital.location = location
        hospital.address = address
        hospital.save()
        return redirect('../HospitalLogin')
    else:
        return render(request,'./Equipment/login.html')

def HospitalLogin(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        hospital = Hospitals.objects.get(username=username)
        if hospital is None:
            return HttpResponse("Username does not exist")
        if password == hospital.password:
            request.session['hospital'] = hospital.username
            return redirect('../HospitalLanding')
        else:
            return render(request,'./Equipment/login.html',{'error':"Invalid Password"})
    else:
        return render(request,'./Equipment/login.html')

def HospitalAddEquipment(request):
    if request.method == 'POST':
        hospital = request.session['hospital']
        equipmentname = request.POST.get('equipmentname')
        equipmentCount = Equipments.objects.filter(name=equipmentname).count()
        if equipmentCount == 0:
            equipment = Equipments()
            equipment.name = equipmentname
            equipment.save()
        if (Equipmentsavailable.objects.filter(equipment=Equipments.objects.get(name=equipmentname), hospital=Hospitals.objects.get(username=hospital)).count()==0):
            equipmentsavailable = Equipmentsavailable()
            equipmentsavailable.hospital = Hospitals.objects.get(username=hospital)
            equipmentsavailable.equipment = Equipments.objects.get(name=equipmentname)
            equipmentsavailable.available = 1
            equipmentsavailable.save()
        else:
            Equipmentsavailable.objects.filter(equipment=Equipments.objects.get(name=equipmentname),hospital=Hospitals.objects.get(username=hospital)).update(available=1)
        return redirect('../HospitalLanding/')
    else:
        return render(request,'./Equipment/AddEquipment.html')

def HospitalLanding(request):
    if request.method =='POST':
        pass
    else:
        hospital = request.session['hospital']
        equipmentsavailable = Equipmentsavailable.objects.all()
        return render(request,'./Equipment/EquipmentTable.html',{'equipments': equipmentsavailable})

def HospitalBookEquipment(request):
    hospital = request.session['hospital']
    equipmentname = request.POST.get('equipmentname')

    if (Equipmentsavailable.objects.filter(equipment=Equipments.objects.get(name=equipmentname)).filter(hospital=Hospitals.objects.get(username=hospital)).count()==0):
        equipmentsavailable = Equipmentsavailable()
        equipmentsavailable.hospital = Hospitals.objects.get(username=hospital)
        equipmentsavailable.equipment = Equipments.objects.get(name=equipmentname)
        equipmentsavailable.available = 1
        equipmentsavailable.save()
    elif (Equipmentsavailable.objects.filter(equipment=Equipments.objects.get(name=equipmentname)).filter(hospital=Hospitals.objects.get(username=hospital))[0].available==1):
        equipmentsavailable = Equipmentsavailable.objects.all()
        return render(request,'./Equipment/EquipmentTable.html',{'error':"Hospital already has the required Equipment Available", 'equipments': equipmentsavailable})
    else:
        Medicinestocks.objects.filter(equipment=Equipments.objects.get(name=medicinename)).filter(hospital=Hospitals.objects.get(username=hospital)).update(available=1)
    return redirect('../HospitalLanding')

def DoctorCheckIn(request):
    if request.method == 'POST':
        hospital = Hospitals.objects.get(username=request.session['hospital'])
        username=request.POST.get('doctor')
        Doctors.objects.filter(id=username).update(currentlocation=hospital)
        Doctors.objects.filter(id=username).update(available=1)
        return render(request,'./qr.html',{'message':'Doctor Checked in successfully'})
    else:
        return render(request,'./qr.html')

def DoctorCheckOut(request):
    if request.method == 'POST':
        hospital = Hospitals.objects.get(username= -1)
        username=request.POST.get('doctor')
        Doctors.objects.filter(id=username).update(currentlocation= -1)
        Doctors.objects.filter(id=username).update(available=0)
        return render(request,'./checkout.html',{'message':'Doctor Checked out successfully'})
    else:
        return render(request,'./checkout.html')


'''
Doctor: videocall requests list

Pharmacy: generate invoices

Hospitals: qr code recognizer(plugin if poss)

API: new api to fetch the pharmacy for a particular medicine.
'''
